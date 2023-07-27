from functions import *
from docx import Document


if __name__ == "__main__":
    # variable to hold chinese words and their translations
    chinese_words = set()
    dictionary = {}

    for each in get_files():
        chinese = find_chinese(read_file(each))
        chinese_words.update(chinese)
    chinese_words = list(chinese_words)

    chinese_doc = Document()

    for each in chinese_words:
        chinese_doc.add_paragraph(each)

    chinese_doc.save("res.docx")

    input(
        "the output file is saved as res.docx open the google translate and paste it here as out.docx.\nonce completed press any key to continue..."
    )

    english_doc = Document("./out.docx")

    if len(chinese_doc.paragraphs) == len(english_doc.paragraphs):
        print("Invalid document, the length of translations doesn't match")

    for i, j in zip(chinese_doc.paragraphs, english_doc.paragraphs):
        dictionary[i] = j

    for file in get_files():
        content = read_file(file)
        words = list(find_chinese(content))
        words.sort(key=lambda x: len(x), reverse=True)

        for word in words:
            content = content.replace(word, dictionary[word])

        # write_file(file.replace(r".\src\amis", r".\src\res"), content)
        write_file(file, content)
