from docx import Document
import csv

out = Document("out.docx")
inp = Document("res.docx")

if len(out.paragraphs) != len(inp.paragraphs):
    print("the lines don't match")

fields = ['Chinese', 'English'] 
filename = "dictionary.csv"
    
with open(filename, 'w', encoding="utf8") as csvfile: 
    csvwriter = csv.writer(csvfile) 
        
    csvwriter.writerow(fields)

    a = inp.paragraphs
    b = out.paragraphs

    count = 0
    for i,j in zip(a, b):
        csvwriter.writerow([i.text, j.text])
        count += 1
        print(f"\r{count}          ", end="")
